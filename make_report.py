"""
make_report.py
--------------
Generates the final summary Word document for the reproducible package.
Run from the package root:  python make_report.py
Output: Career_Outcome_Visualization_Report.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── colour constants ──────────────────────────────────────────────────────────
NAVY       = RGBColor(0x12, 0x2A, 0x52)   # heading 1
DARK_BLUE  = RGBColor(0x2F, 0x6D, 0xB3)   # heading 2 / accent
MED_BLUE   = RGBColor(0x4A, 0x84, 0xC4)   # heading 3
GRAY_TEXT  = RGBColor(0x44, 0x44, 0x44)   # body text
LIGHT_GRAY = RGBColor(0xF2, 0xF2, 0xF2)   # table shading

OUTPUT_PATH  = Path(__file__).parent / "Career_Outcome_Visualization_Report.docx"
PACKAGE_DIR  = Path(__file__).parent


# ── helpers ───────────────────────────────────────────────────────────────────

def shade_cell(cell, hex_rgb: str) -> None:
    """Apply background shading to a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_rgb)
    tcPr.append(shd)


def add_h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text.upper())
    run.bold      = True
    run.font.size = Pt(16)
    run.font.color.rgb = NAVY
    # bottom border
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "8")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "2F6DB3")
    pBdr.append(bot)
    pPr.append(pBdr)


def add_h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(13)
    run.font.color.rgb = DARK_BLUE


def add_h3(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(11)
    run.font.color.rgb = MED_BLUE


def add_body(doc: Document, text: str, italic: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = GRAY_TEXT
    run.italic = italic


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent   = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_before  = Pt(1)
    p.paragraph_format.space_after   = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = GRAY_TEXT


def add_note(doc: Document, text: str) -> None:
    """Indented italic note box."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run("ℹ  " + text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_two_col_table(doc: Document, rows: list[tuple[str, str]],
                      header: tuple[str, str] | None = None,
                      col_widths: tuple[float, float] = (2.5, 4.0)) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    if header:
        hrow = table.add_row()
        for i, h in enumerate(header):
            cell = hrow.cells[i]
            cell.width = Inches(col_widths[i])
            shade_cell(cell, "2F6DB3")
            run = cell.paragraphs[0].add_run(h)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for idx, (a, b) in enumerate(rows):
        row = table.add_row()
        row.cells[0].width = Inches(col_widths[0])
        row.cells[1].width = Inches(col_widths[1])
        if idx % 2 == 0:
            shade_cell(row.cells[0], "F2F4F8")
            shade_cell(row.cells[1], "F2F4F8")
        ra = row.cells[0].paragraphs[0].add_run(a)
        rb = row.cells[1].paragraphs[0].add_run(b)
        ra.font.size = rb.font.size = Pt(10)
        ra.font.color.rgb = rb.font.color.rgb = GRAY_TEXT

    doc.add_paragraph()  # spacing after table


def add_three_col_table(doc: Document, rows: list[tuple[str, str, str]],
                        header: tuple[str, str, str] | None = None,
                        col_widths: tuple[float, float, float] = (1.5, 2.2, 2.8)) -> None:
    table = doc.add_table(rows=0, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    if header:
        hrow = table.add_row()
        for i, h in enumerate(header):
            cell = hrow.cells[i]
            shade_cell(cell, "2F6DB3")
            run = cell.paragraphs[0].add_run(h)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for idx, (a, b, c) in enumerate(rows):
        row = table.add_row()
        if idx % 2 == 0:
            for cell in row.cells:
                shade_cell(cell, "F2F4F8")
        for cell, val in zip(row.cells, (a, b, c)):
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)
            run.font.color.rgb = GRAY_TEXT

    doc.add_paragraph()


# ── cover page ────────────────────────────────────────────────────────────────

def build_cover(doc: Document) -> None:
    # top spacing
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Career Outcome Visualization")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Reproducible Package — Final Summary Report")
    r2.font.size = Pt(16)
    r2.font.color.rgb = DARK_BLUE

    doc.add_paragraph()

    info_lines = [
        ("Project", "Gonzaga University Career Services — Major to Career Pathway Analysis"),
        ("Cohort", "Class of 2026"),
        ("Dashboards", "3 interactive Circos chord-diagram visualizations"),
        ("Prepared by", "Graduate Research Project"),
        ("Date", datetime.date.today().strftime("%B %d, %Y")),
    ]
    for label, value in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lb = p.add_run(f"{label}:  ")
        lb.bold = True
        lb.font.size = Pt(11)
        lb.font.color.rgb = DARK_BLUE
        vl = p.add_run(value)
        vl.font.size = Pt(11)
        vl.font.color.rgb = GRAY_TEXT

    doc.add_page_break()


# ── main document ─────────────────────────────────────────────────────────────

def build_document() -> None:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Default font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    build_cover(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # 1. EXECUTIVE SUMMARY
    # ─────────────────────────────────────────────────────────────────────────
    add_h1(doc, "1. Executive Summary")
    add_body(doc,
        "This report documents the Career Outcome Visualization project developed over four months "
        "(January–May 2026) as part of a graduate research engagement with Gonzaga University's "
        "Career & Professional Development office. The project produced three interactive, "
        "browser-based chord-diagram dashboards that visualize how graduating students move from "
        "their incoming academic major to their career outcome or graduating major."
    )
    add_body(doc,
        "All three dashboards are fully functional static HTML files that can be opened in any "
        "modern browser without a server, database, or internet connection. The reproducible "
        "package collects every file, script, and decision record needed for a new student or "
        "staff member to view, understand, and regenerate the dashboards from fresh data."
    )

    add_h2(doc, "Three Deliverables")
    add_two_col_table(doc, [
        ("Career Outcomes Dashboard\nsite/index.html",
         "Shows how 2,682 graduates mapped from their incoming major to one of 8 career clusters. "
         "Clicking any major or career thumbnail filters the central chord diagram."),
        ("Industry Mobility Dashboard\nsite/incoming_grad_grouped.html",
         "Shows how 925 students moved from their incoming major group to their graduating major "
         "group, organized by 14 industry-style groups (e.g., Computing/Data, Engineering, Education)."),
        ("College Mobility Dashboard\nsite/incoming_grad_grouped_college.html",
         "Same 925 students viewed by college (OCB, CAS, EDUC, COH, SOE, Other) — shows "
         "cross-college major switching patterns."),
    ], header=("Dashboard", "Description"), col_widths=(2.4, 4.1))

    doc.add_page_break()

    # ── Folder tree diagram ──────────────────────────────────────────────────
    add_h1(doc, "Package Folder Structure")
    add_body(doc,
        "The diagram below shows every file and folder in the reproducible package. "
        "Orange badges are folders; blue badges are files. Annotations on the right "
        "describe each item's purpose."
    )
    tree_img = PACKAGE_DIR / "folder_tree.png"
    if tree_img.exists():
        doc.add_picture(str(tree_img), width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # 2. PROJECT BACKGROUND
    # ─────────────────────────────────────────────────────────────────────────
    add_h1(doc, "2. Project Background")
    add_body(doc,
        "Gonzaga University's Career & Professional Development office collects annual career "
        "outcome survey data from graduating students. The raw data contains each graduate's "
        "declared major and self-reported job title. A parallel dataset from the Registrar tracks "
        "students' declared major at entry (CPF semester) versus at graduation (SP26 semester)."
    )
    add_body(doc,
        "Prior to this project, no visual tool existed that allowed career advisors or students "
        "to explore the relationship between academic major and career outcome at a glance. "
        "The chord-diagram (Circos) format was chosen because it shows both the volume of "
        "connections and the directionality between two sets of categories simultaneously."
    )

    add_h2(doc, "Why Circos?")
    add_body(doc,
        "Circos is an open-source, Perl-based command-line tool developed for genomics research "
        "that produces publication-quality chord diagrams. It was chosen over D3.js or Plotly "
        "for three reasons:"
    )
    add_bullet(doc, "Precise geometric control — ribbon shapes, arc spacing, label positions, "
               "and color assignments can all be specified exactly in configuration files.")
    add_bullet(doc, "Static output — Circos renders a single SVG file per diagram. These SVGs "
               "are embedded directly in HTML pages, making the dashboards fast, printable, "
               "and usable offline.")
    add_bullet(doc, "Separation of data and presentation — Circos configuration files are "
               "plain-text, making it straightforward to update the data and re-render.")

    add_h2(doc, "Development Timeline")
    add_two_col_table(doc, [
        ("January 2026",   "Initial data exploration; built first prototype Circos diagrams from original lookup table."),
        ("February 2026",  "Identified severe data loss issue (only ~1,100 of 4,129 graduates classified). Rebuilt cluster lookup pipeline from scratch."),
        ("March 2026",     "Implemented rule-based job-title heuristics; raised classified count to 2,682. Rebuilt major and career cluster consolidation (Steps 2–3)."),
        ("April 2026",     "Ran 7 layout experiments (v1–v7) to optimize arc ordering and minimize ribbon crossings. Finalized v3 ordering."),
        ("Early May 2026", "Developed SVG post-processing (taper, font replacement). Added industry and college mobility dashboards."),
        ("Mid May 2026",   "Finalized dashboard HTML, documentation panel, and reproducible package."),
    ], header=("Period", "Milestones"), col_widths=(1.6, 4.9))

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # 3. DATA SOURCES AND INPUT FILES
    # ─────────────────────────────────────────────────────────────────────────
    add_h1(doc, "3. Data Sources and Input Files")

    add_h2(doc, "Career Survey Data")
    add_body(doc,
        "File: career_survey_data.xlsx  |  4,129 rows  |  Source: Career & Professional Development office"
    )
    add_bullet(doc, "Each row represents one graduate response from the annual career outcomes survey.")
    add_bullet(doc, "Key columns used: Program Name/Major (the student's declared major at graduation), "
               "Job Title (self-reported first destination career role).")
    add_bullet(doc, "1,338 rows had a blank Job Title field — these students did not report a "
               "career outcome and are excluded from the career dashboard.")
    add_bullet(doc, "After full classification, 2,682 graduates appear in the career dashboard.")

    add_h2(doc, "Enrollment Transition Data")
    add_body(doc,
        "File: site/incomingmajor_fall2022_csv/Class_of_2026_CPF_vs_SP26.csv  |  925 rows"
    )
    add_bullet(doc, "Each row represents one student with both a recorded incoming major (CPF "
               "semester = first full-time semester) and a recorded graduating major (SP26 = "
               "Spring 2026 graduation semester).")
    add_bullet(doc, "These 925 students form the basis for both mobility dashboards.")
    add_bullet(doc, "Students who transferred in late, left the university, or had no recorded "
               "incoming major are not included.")

    add_h2(doc, "Major Grouping File")
    add_body(doc,
        "File: site/incomingmajor_fall2022_csv/major.csv  |  Created manually by project team"
    )
    add_bullet(doc, "Maps each individual major name to a display group (e.g., Computer Science "
               "→ CS/DS/DA/Math/Stat) and a college abbreviation (e.g., CAS, OCB, COH).")
    add_bullet(doc, "Used exclusively by the two mobility dashboards (Steps 7 and 9).")

    add_h2(doc, "Baseline Reference Files (inputs/ folder)")
    add_two_col_table(doc, [
        ("Major_Job_Cluster_Lookup_Rebuilt_Full_baseline.csv",
         "Row-level lookup of every graduate's major + job title after initial heuristic "
         "classification. Input to the major and career cluster consolidation steps."),
        ("Major_Job_Cluster_Lookup_Filtered_baseline.csv",
         "Earlier version of the filtered lookup (kept for reference and comparison)."),
        ("job_title_categorized.csv",
         "~87 manually reviewed job titles with assigned categories. Overrides heuristic "
         "classification for ambiguous titles like 'Associate', 'Analyst', 'Consultant'."),
        ("Major_Career_Analysis_v4__Cluster_Breakdown.csv",
         "External major cluster reference from prior analysis project. Input to Step 1."),
        ("Major_Career_Analysis_v4__Career_Analysis.csv",
         "External career cluster reference from prior analysis project. Input to Step 1."),
    ], header=("File", "Purpose"), col_widths=(2.9, 3.6))

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # 4. PIPELINE OVERVIEW
    # ─────────────────────────────────────────────────────────────────────────
    add_h1(doc, "4. Pipeline Overview — 10 Steps")
    add_body(doc,
        "The pipeline converts raw Excel/CSV data into pre-rendered SVG images embedded in the "
        "three dashboard HTML files. It is divided into three tracks:"
    )
    add_bullet(doc, "Career track (Steps 1–6): raw survey → classified lookup → Circos config → SVG → career dashboard")
    add_bullet(doc, "Industry mobility track (Steps 7–8): enrollment CSV → Circos config → SVG → industry dashboard")
    add_bullet(doc, "College mobility track (Steps 9–10): same enrollment CSV → Circos config → SVG → college dashboard")
    add_body(doc, "Run all steps at once with:  python run_pipeline.py", italic=True)

    add_h2(doc, "Step-by-Step Reference")
    add_three_col_table(doc, [
        ("Step 01\n01_rebuild_cluster_lookup.py",
         "career_survey_data.xlsx +\nexternal cluster CSVs",
         "Reads every graduate row from Excel. Classifies each major using a reference table "
         "plus alias overrides. Classifies each job title using exact match → normalized match "
         "→ keyword heuristic rules. Outputs a full row-level lookup and a filtered version "
         "(rows with both major and career classified, non-blank job title)."),
        ("Step 02\n02_remap_major_clusters.py",
         "inputs/Major_Job_Cluster_Lookup_\nRebuilt_Full_baseline.csv",
         "Consolidates 24 raw major cluster names into 13 display clusters grouped under "
         "4 large academic families: Engineering/Science, Business, Humanities, Social Sciences."),
        ("Step 03\n03_remap_career_clusters.py",
         "outputs/lookup/\nMajor_Job_Cluster_Lookup_v2_full.csv",
         "Consolidates 14 raw career clusters into 8 display clusters. Applies manual overrides "
         "from job_title_categorized.csv before fallback rules. Produces the final production "
         "lookup: Major_Job_Cluster_Lookup_v2_filtered_career_merged.csv."),
        ("Step 04\n04_generate_circos_data_career.py",
         "outputs/lookup/\n...v2_filtered_career_merged.csv",
         "Converts the lookup CSV into Circos input format: karyotype.txt (arc sizes and "
         "colors), links.txt (ribbon connections), circos.conf (master config), colors.conf. "
         "Applies the v3 arc ordering to minimize ribbon crossings."),
        ("Step 05\n05_taper_ribbon_svg.py",
         "Raw SVG from Circos\n(Step 04 output)",
         "Post-processes the Circos SVG to taper ribbon paths — ribbons narrow to 28% of "
         "their endpoint width at the midpoint, then widen again. Also replaces Circos bitmap "
         "font references with Segoe UI system fonts and applies bold weight to arc labels."),
        ("Step 06\n06_generate_career_focus_svgs.py",
         "Tapered SVG from Step 05",
         "Produces 22 SVG files: one overall diagram, 13 per-major-cluster views, and 8 "
         "per-career-cluster views. Each focus SVG is the full diagram with all ribbons except "
         "the selected cluster's ribbons set to transparent. Copies to site/v3_ordered_images/."),
        ("Step 07\n07_generate_circos_data_industry.py",
         "Class_of_2026_CPF_vs_SP26.csv\n+ major.csv",
         "Groups incoming and graduating majors into 14 industry-style display groups. "
         "Counts student transitions between each pair of groups. Generates Circos config "
         "files in outputs/renders/incoming_grad_grouped_latest/."),
        ("Step 08\n08_generate_industry_focus_svgs.py",
         "Main SVG from Step 07",
         "Produces 27 SVG files: overall + 13 incoming group views + 13 graduating group "
         "views. Copies to site/incoming_grad_grouped_dashboard_images/."),
        ("Step 09\n09_generate_circos_data_college.py",
         "Class_of_2026_CPF_vs_SP26.csv\n+ major.csv",
         "Same logic as Step 07 but groups students by college (OCB, CAS, EDUC, COH, SOE, "
         "Other) rather than industry group. Generates config in "
         "outputs/renders/incoming_grad_grouped_college_latest/."),
        ("Step 10\n10_generate_college_focus_svgs.py",
         "Main SVG from Step 09",
         "Produces 38 SVG files: overall + per-college incoming/graduating views + "
         "per-group views. Copies to site/incoming_grad_grouped_college_dashboard_images/."),
    ], header=("Script", "Input", "What it does"), col_widths=(1.6, 1.7, 3.2))

    # ── Pipeline flow diagram ────────────────────────────────────────────────
    add_h2(doc, "Visual Pipeline Map")
    add_body(doc,
        "The diagram below shows the complete data flow from raw input files through "
        "all 10 pipeline steps to the three HTML dashboards."
    )
    flow_img = PACKAGE_DIR / "pipeline_flow.png"
    if flow_img.exists():
        doc.add_picture(str(flow_img), width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # 5. DATA CLEANING AND CLUSTERING DECISIONS
    # ─────────────────────────────────────────────────────────────────────────
    add_h1(doc, "5. Data Cleaning and Clustering Decisions")
    add_body(doc,
        "Every significant decision in the data pipeline is documented below. These decisions "
        "were made collaboratively during the project and represent the rationale that should "
        "be reviewed when new cohort data arrives."
    )

    add_h2(doc, "Decision 1 — What counts as a classifiable record?")
    add_body(doc,
        "A graduate's record is included in the career dashboard if and only if: "
        "(a) their major matched a cluster, AND (b) their job title field is not blank, "
        "AND (c) the job title was classified into a non-Other career cluster."
    )
    add_bullet(doc, "1,338 of 4,129 rows had no job title → excluded (cannot infer career from blank).")
    add_bullet(doc, "~89 rows had job titles that could not be classified after all heuristics → "
               "excluded from dashboard but retained in the full lookup CSV for reference.")
    add_bullet(doc, "Result: 2,682 graduates appear on the career dashboard out of 4,129 total responses.")

    add_h2(doc, "Decision 2 — Why rebuild the lookup from scratch?")
    add_body(doc,
        "The original pre-existing lookup table covered only ~1,078 exact job title matches out of "
        "thousands of raw titles, producing only ~1,100 classified records. A three-tier "
        "classification was built: exact match → normalized match (strip punctuation/case) → "
        "keyword heuristic rules. This raised classified records to 2,682."
    )
    add_note(doc, "Risk: keyword heuristics can misclassify edge cases (e.g., 'research analyst' "
             "at a bank classified as Research & Science instead of Finance). This was an accepted "
             "trade-off: coverage was prioritized over precision for a public aggregate visualization.")

    add_h2(doc, "Decision 3 — Economics is Social Science, not Business")
    add_body(doc,
        "All Economics variants (Economics, Economics - Business, Economics - International, "
        "Economics - Mathematical, Economics - Public Policy) were placed in "
        "SOCIAL SCIENCES & EDUCATION, not in BUSINESS & MANAGEMENT."
    )
    add_bullet(doc, "Gonzaga Economics is taught in the College of Arts & Sciences, not the School of Business.")
    add_bullet(doc, "US News and Carnegie Classification both categorize Economics as Social Science.")
    add_bullet(doc, "Merging Economics into Business would inflate the Business arc artificially.")

    add_h2(doc, "Decision 4 — Major cluster consolidation (24 → 13 display clusters)")
    add_three_col_table(doc, [
        ("Engineering Disciplines",                    "Engineering",                       "ENGINEERING, SCIENCE & TECH"),
        ("Computer Science & IT",                      "Computer Science & IT",             "ENGINEERING, SCIENCE & TECH"),
        ("Biological Sciences + Health & Wellness",    "Biological & Health Sciences",      "ENGINEERING, SCIENCE & TECH"),
        ("Physical Sciences + Environmental Sciences", "Physical & Environmental Sciences", "ENGINEERING, SCIENCE & TECH"),
        ("Finance & Accounting",                       "Finance & Accounting",              "BUSINESS & MANAGEMENT"),
        ("General Business",                           "General Business",                  "BUSINESS & MANAGEMENT"),
        ("Operations & Analytics",                     "Operations & Analytics",            "BUSINESS & MANAGEMENT"),
        ("Specialized Business",                       "Specialized Business",              "BUSINESS & MANAGEMENT"),
        ("Marketing + Communication + Journalism + Creative Writing", "Marketing & Communication", "HUMANITIES & COMMUNICATION"),
        ("Arts & Theology + English & Writing + Languages", "Languages, Arts & Humanities", "HUMANITIES & COMMUNICATION"),
        ("Teacher Ed + Social Work + Spec Ed + Ed Leadership", "Education & Social Services", "SOCIAL SCIENCES & EDUCATION"),
        ("Social Sciences + Humanities & Liberal Arts + Intl Studies", "Social Sciences & Humanities", "SOCIAL SCIENCES & EDUCATION"),
        ("Economics",                                  "Economics",                         "SOCIAL SCIENCES & EDUCATION"),
    ], header=("Raw cluster(s) merged", "Display cluster", "Large family"), col_widths=(2.2, 2.1, 2.2))

    add_h2(doc, "Decision 5 — Career cluster consolidation (14 → 8 display clusters)")
    add_two_col_table(doc, [
        ("Finance & Investment + Accounting & Audit",    "Finance & Accounting"),
        ("Management & Operations + Leadership (General)", "Management & Ops"),
        ("Software & Data + IT & Infrastructure",        "Software & Data"),
        ("Clinical Care + Research & Science",           "Health & Science Research"),
        ("Education + Social Service",                   "Education & Social Service"),
        ("Sales & Marketing",                            "Sales & Marketing (unchanged)"),
        ("Legal & Policy",                               "Legal & Policy (unchanged)"),
        ("Engineering",                                  "Engineering (unchanged)"),
    ], header=("Raw career cluster(s)", "Display cluster"), col_widths=(3.3, 3.2))

    add_h2(doc, "Decision 6 — Manual job title overrides (87 titles)")
    add_body(doc,
        "Titles that keyword heuristics could not reliably place were manually reviewed and "
        "assigned a category in inputs/job_title_categorized.csv. Examples:"
    )
    add_two_col_table(doc, [
        ("Associate",           "Management & Ops — too generic; most entry-level associates are in ops roles"),
        ("Analyst",             "Software & Data — generic; treated as data given university's CS graduate count"),
        ("Business Analyst",    "Management & Ops — process improvement focus, distinct from data analyst"),
        ("Scribe",              "Health & Science Research — medical scribe = clinical support"),
        ("RN",                  "Health & Science Research — Registered Nurse abbreviation"),
        ("Server / Bartender",  "Management & Ops — hospitality/food service, operational role"),
        ("Consultant",          "Management & Ops — general consulting, not IT-specific"),
        ("Missionary",          "Education & Social Service — humanitarian/service role"),
    ], header=("Job title", "Assignment and rationale"), col_widths=(1.8, 4.7))

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # 6. VISUALIZATION DESIGN RULES
    # ─────────────────────────────────────────────────────────────────────────
    add_h1(doc, "6. Visualization Design Rules")

    add_h2(doc, "Color Palette")
    add_body(doc,
        "The palette was designed to be colorblind-friendly (no red-green pairing, which affects "
        "~8% of men) while maintaining clear contrast between the four academic families. "
        "Career arcs use neutral gray so the eye reads the colored major arcs as the 'source'."
    )
    add_two_col_table(doc, [
        ("BUSINESS & MANAGEMENT",             "#2F6DB3 — medium blue (shaded lighter for sub-clusters)"),
        ("ENGINEERING, SCIENCE & TECHNOLOGY", "#2E8B57 — deep teal-green (shaded lighter)"),
        ("SOCIAL SCIENCES & EDUCATION",       "#E68613 — vivid orange (shaded lighter)"),
        ("HUMANITIES & COMMUNICATION",        "#B279A2 — soft purple-magenta (shaded lighter)"),
        ("Career arcs (all clusters)",         "#B9B9B9 — neutral medium gray"),
    ], header=("Academic family", "Color"), col_widths=(3.1, 3.4))

    add_h2(doc, "Arc Ordering (clockwise from top)")
    add_body(doc,
        "Major clusters are ordered to minimize ribbon crossings. Related clusters that share "
        "graduates with the same career destinations are placed adjacent to each other."
    )
    add_two_col_table(doc, [
        ("Engineering",                    "ENGINEERING, SCIENCE & TECH"),
        ("Computer Science & IT",          "ENGINEERING, SCIENCE & TECH"),
        ("Biological & Health Sciences",   "ENGINEERING, SCIENCE & TECH"),
        ("Physical & Environmental Sci.",  "ENGINEERING, SCIENCE & TECH"),
        ("Finance & Accounting",           "BUSINESS & MANAGEMENT"),
        ("General Business",               "BUSINESS & MANAGEMENT"),
        ("Operations & Analytics",         "BUSINESS & MANAGEMENT"),
        ("Specialized Business",           "BUSINESS & MANAGEMENT"),
        ("Marketing & Communication",      "HUMANITIES & COMMUNICATION"),
        ("Languages, Arts & Humanities",   "HUMANITIES & COMMUNICATION"),
        ("Education & Social Services",    "SOCIAL SCIENCES & EDUCATION"),
        ("Social Sciences & Humanities",   "SOCIAL SCIENCES & EDUCATION"),
        ("Economics",                      "SOCIAL SCIENCES & EDUCATION"),
    ], header=("Major cluster (clockwise order)", "Large family"), col_widths=(3.1, 3.4))

    add_h2(doc, "SVG Ribbon Tapering")
    add_body(doc,
        "Circos draws ribbons as uniform-width bezier curves. A custom SVG post-processing "
        "script (Step 05) reshapes each ribbon to taper toward the center of the diagram "
        "and widen at both arc attachment points. This is a standard practice in scientific "
        "chord diagrams that improves readability when many ribbons overlap."
    )
    add_bullet(doc, "Mid-scale: 0.28 — ribbon width at midpoint is 28% of endpoint width.")
    add_bullet(doc, "Profile power: 1.35 — controls how quickly the taper narrows toward the center.")
    add_bullet(doc, "Font replacement: Circos default fonts are replaced with Segoe UI / Tahoma / Verdana.")

    add_h2(doc, "Pre-Rendered Focus SVGs")
    add_body(doc,
        "The 'click a thumbnail to filter the diagram' interaction is achieved by pre-rendering "
        "a separate SVG for each cluster — the full diagram where all other clusters' ribbons "
        "are transparent. The browser simply swaps the src attribute of the center image tag. "
        "This approach loads instantly and works without a server."
    )
    add_two_col_table(doc, [
        ("site/v3_ordered_images/",                          "22 SVGs — 1 overall + 13 major + 8 career"),
        ("site/incoming_grad_grouped_dashboard_images/",     "27 SVGs — 1 overall + 13 incoming + 13 graduating"),
        ("site/incoming_grad_grouped_college_dashboard_images/", "38 SVGs — 1 overall + college + group views"),
    ], header=("Folder", "Contents"), col_widths=(3.6, 2.9))

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # 7. PACKAGE FILE INVENTORY
    # ─────────────────────────────────────────────────────────────────────────
    add_h1(doc, "7. Package File Inventory")
    add_body(doc,
        "The complete file set needed to view the dashboards, understand every decision, "
        "and regenerate from new data. Everything here is in the reproducible_package/ folder."
    )

    add_h2(doc, "Documentation files")
    add_two_col_table(doc, [
        ("README.md",                    "Quick start + package guide + how to update for a new cohort"),
        ("PIPELINE_OVERVIEW.md",         "Full data flow diagram + per-step explanation + common errors"),
        ("DATA_CLEANING_DECISIONS.md",   "10 documented data decisions with rationale"),
        ("VISUALIZATION_RULES.md",       "Color system, arc ordering, taper params, SVG structure"),
        ("CIRCOS_SETUP.md",              "Circos installation steps for Windows, macOS, and Linux"),
        ("requirements.txt",             "Python dependencies: pandas, openpyxl"),
        ("run_pipeline.py",              "Master runner — runs all 10 steps in order"),
        ("Career_Outcome_Visualization_Report.docx", "This document"),
    ], header=("File", "Contents"), col_widths=(3.2, 3.3))

    add_h2(doc, "Input data files")
    add_two_col_table(doc, [
        ("career_survey_data.xlsx",                             "Raw career survey (4,129 rows)"),
        ("inputs/Major_Job_Cluster_Lookup_Rebuilt_Full_baseline.csv", "Baseline lookup (all graduates classified)"),
        ("inputs/Major_Job_Cluster_Lookup_Filtered_baseline.csv",     "Earlier filtered lookup (reference)"),
        ("inputs/job_title_categorized.csv",                    "87 manual job-title category overrides"),
        ("inputs/Major_Career_Analysis_v4__Cluster_Breakdown.csv",    "External major cluster reference"),
        ("inputs/Major_Career_Analysis_v4__Career_Analysis.csv",      "External career cluster reference"),
        ("site/incomingmajor_fall2022_csv/Class_of_2026_CPF_vs_SP26.csv", "Enrollment transition data (925 students)"),
        ("site/incomingmajor_fall2022_csv/major.csv",           "Major → group → college mapping"),
    ], header=("File", "Description"), col_widths=(3.8, 2.7))

    add_h2(doc, "Pre-built outputs (ready to use without running the pipeline)")
    add_two_col_table(doc, [
        ("outputs/lookup/Major_Job_Cluster_Lookup_v2_filtered_career_merged.csv",
         "Final production lookup — 2,682 classified graduates with major and career clusters"),
        ("site/index.html",                                    "Career outcomes dashboard (opens in browser)"),
        ("site/incoming_grad_grouped.html",                    "Industry mobility dashboard"),
        ("site/incoming_grad_grouped_college.html",            "College mobility dashboard"),
        ("site/dashboard_docs.css / .js",                      "Shared stylesheet and interaction JavaScript"),
        ("site/dashboard_docs_data.js",                        "Methodology text for docs panel (hand-maintained)"),
        ("site/v3_ordered_images/ (22 SVGs)",                  "Career dashboard graphics"),
        ("site/incoming_grad_grouped_dashboard_images/ (27 SVGs)", "Industry dashboard graphics"),
        ("site/incoming_grad_grouped_college_dashboard_images/ (38 SVGs)", "College dashboard graphics"),
    ], header=("File / folder", "Description"), col_widths=(3.8, 2.7))

    add_h2(doc, "Pipeline scripts (scripts/ folder)")
    add_two_col_table(doc, [
        ("01_rebuild_cluster_lookup.py",        "Excel → row-level classified lookup (Steps 1 of career track)"),
        ("02_remap_major_clusters.py",           "Consolidate major clusters (24 → 13)"),
        ("03_remap_career_clusters.py",          "Consolidate career clusters (14 → 8), produce final lookup"),
        ("04_generate_circos_data_career.py",    "Lookup CSV → Circos config files for career dashboard"),
        ("05_taper_ribbon_svg.py",               "Post-process SVG: taper ribbons + replace fonts"),
        ("06_generate_career_focus_svgs.py",     "Extract 22 per-cluster focus SVGs for career dashboard"),
        ("07_generate_circos_data_industry.py",  "Enrollment CSV → Circos config for industry mobility"),
        ("08_generate_industry_focus_svgs.py",   "Extract 27 focus SVGs for industry dashboard"),
        ("09_generate_circos_data_college.py",   "Enrollment CSV → Circos config for college mobility"),
        ("10_generate_college_focus_svgs.py",    "Extract 38 focus SVGs for college dashboard"),
    ], header=("Script", "Role"), col_widths=(3.0, 3.5))

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # 8. HOW TO USE THIS PACKAGE
    # ─────────────────────────────────────────────────────────────────────────
    add_h1(doc, "8. How to Use This Package")

    add_h2(doc, "Option A — View the Existing Dashboards (No Installation Needed)")
    add_body(doc, "No Python. No Circos. No server. Just open one of these files in a browser:")
    add_bullet(doc, "site/index.html — Career Outcomes Dashboard")
    add_bullet(doc, "site/incoming_grad_grouped.html — Industry Mobility Dashboard")
    add_bullet(doc, "site/incoming_grad_grouped_college.html — College Mobility Dashboard")
    add_note(doc, "Tip: Right-click the file → Open with → your browser. "
             "Do not open from a restricted network share or Windows Explorer preview pane.")

    add_h2(doc, "Option B — Regenerate from Existing Data (Partial Pipeline)")
    add_body(doc,
        "If you want to re-render the SVGs without changing the underlying data "
        "(e.g., to adjust colors or spacing):"
    )
    add_bullet(doc, "Step 1: Install Circos — see CIRCOS_SETUP.md.")
    add_bullet(doc, "Step 2: Set CIRCOS_ROOT environment variable to your Circos install folder.")
    add_bullet(doc, "Step 3: Run only Steps 4–6 for the career dashboard, or Steps 7–10 for the mobility dashboards:")
    add_body(doc, "    python run_pipeline.py --steps 4 5 6", italic=True)

    add_h2(doc, "Option C — Update with New Cohort Data (Full Pipeline)")
    add_body(doc, "When new data arrives (e.g., Class of 2027):")
    add_bullet(doc, "Replace career_survey_data.xlsx with the new Excel export. "
               "Verify columns: Program Name/Major, Job Title.")
    add_bullet(doc, "Replace site/incomingmajor_fall2022_csv/Class_of_2026_CPF_vs_SP26.csv "
               "with the new enrollment transition file. Verify column names match.")
    add_bullet(doc, "Update site/incomingmajor_fall2022_csv/major.csv if any new majors were added.")
    add_bullet(doc, "Install Python dependencies: pip install -r requirements.txt")
    add_bullet(doc, "Run the full pipeline: python run_pipeline.py")
    add_bullet(doc, "Check outputs/lookup/Unmatched_Job_Titles_Summary.csv for any new "
               "unclassified job titles. Add high-frequency titles to inputs/job_title_categorized.csv.")
    add_bullet(doc, "Update graduate counts in site/dashboard_docs_data.js.")
    add_bullet(doc, "Update cache-buster query strings in each HTML file "
               "(e.g., change ?v=20260507 to ?v=20270515).")

    add_h2(doc, "Setup Requirements Summary")
    add_two_col_table(doc, [
        ("Python",         "3.10 or newer — https://python.org"),
        ("pandas",         "pip install pandas  (or: pip install -r requirements.txt)"),
        ("openpyxl",       "pip install openpyxl  (needed to read .xlsx files)"),
        ("Circos",         "v0.69-9 — https://circos.ca — see CIRCOS_SETUP.md for full steps"),
        ("Strawberry Perl","Windows only — required to run Circos — https://strawberryperl.com"),
        ("CIRCOS_ROOT",    "Environment variable pointing to your Circos install folder"),
    ], header=("Dependency", "Notes"), col_widths=(1.8, 4.7))

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # 9. KEY METRICS
    # ─────────────────────────────────────────────────────────────────────────
    add_h1(doc, "9. Key Metrics")
    add_two_col_table(doc, [
        ("Total raw career survey rows",          "4,129"),
        ("Rows with blank job title (excluded)",  "1,338"),
        ("Rows classified and shown on dashboard","2,682"),
        ("Major clusters (display level)",        "13"),
        ("Career clusters (display level)",       "8"),
        ("Large academic families",               "4"),
        ("Students in mobility dashboards",       "925"),
        ("Industry groups (mobility)",            "14"),
        ("Colleges (mobility)",                   "6  (OCB, CAS, EDUC, COH, SOE, Other)"),
        ("Total SVGs generated per pipeline run", "87  (22 + 27 + 38)"),
        ("Python scripts in pipeline",            "10"),
        ("Manual job-title overrides",            "~87 titles in job_title_categorized.csv"),
        ("Layout experiments tested",             "7  (v1–v7 arc orderings)"),
        ("Development period",                    "January–May 2026 (approx. 4 months)"),
    ], header=("Metric", "Value"), col_widths=(3.5, 3.0))

    doc.add_page_break()

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # 10. MAJOR MOBILITY ANALYSIS
    # ─────────────────────────────────────────────────────────────────────────
    add_h1(doc, "10. Major Mobility Analysis")
    add_body(doc,
        "The following tables are derived from analysis_outputs/mobility_tables/"
        "mobility_transition_tables.xlsx — a companion analysis file that quantifies "
        "how students moved between major groups and colleges between their incoming "
        "semester and graduation. The dataset covers 925 students from the Class of 2026 "
        "with both a recorded incoming major and a recorded graduating major."
    )
    add_note(doc,
        "How to read the tables: the Incoming row shows where students started. "
        "The Graduating column shows where they ended up. The diagonal (same row and column) "
        "represents students who stayed in the same group. Off-diagonal cells represent "
        "students who switched."
    )

    add_h2(doc, "Major Group Retention Rates")
    add_body(doc,
        "Retention rate = percentage of students who stayed within the same major group "
        "from incoming to graduating semester. Higher retention means the group holds "
        "students well; lower retention means students frequently switch out."
    )
    add_two_col_table(doc, [
        ("Accounting / Finance / Ops",  "80.1%  —  197 of 246 students stayed"),
        ("Engineering",                 "76.6%  —  164 of 214 students stayed"),
        ("Humanities",                  "74.3%  —  52 of 70 students stayed"),
        ("Computing / Data / Math / Stats", "69.7%  —  92 of 132 students stayed"),
        ("Health Sciences",             "68.4%  —  130 of 190 students stayed"),
        ("Marketing",                   "67.2%  —  78 of 116 students stayed"),
        ("Natural Sciences",            "61.4%  —  108 of 176 students stayed"),
        ("Social Sciences",             "59.6%  —  93 of 156 students stayed"),
        ("Writing / Communication",     "58.3%  —  28 of 48 students stayed"),
        ("Education",                   "53.8%  —  14 of 26 students stayed"),
        ("Management",                  "42.6%  —  52 of 122 students stayed"),
        ("Specialized Business",        "33.5%  —  53 of 158 students stayed"),
        ("Pre-Professional / Undeclared", "N/A  —  by definition, all students in this group switch to a declared major"),
    ], header=("Major group", "Retention"), col_widths=(2.8, 3.7))

    add_h2(doc, "Major Group Transition Counts (Incoming → Graduating)")
    add_body(doc,
        "Full transition count matrix. Rows = incoming group. Columns = graduating group. "
        "Row Total = number of students who started in that group."
    )

    # Build the 14-row transition count table
    groups = [
        "Acct/Fin/Ops", "Management", "Spec. Business", "Marketing",
        "Writing/Comm", "Humanities", "Social Sci.", "Pre-Prof/Undecl.",
        "Education", "Health Sci.", "Natural Sci.", "Engineering",
        "Computing/Data", "Col. Total",
    ]
    col_headers = ("Incoming \\ Graduating",
                   "Acct/Fin", "Mgmt", "Spec.Bus", "Mktg",
                   "Writing", "Hum.", "Social", "Pre-Prof",
                   "Educ", "Health", "Nat.Sci", "Eng.", "Comp/Data", "Total")

    # Raw counts from the data
    raw_rows = [
        ("Acct/Fin/Ops (n=246)",    "197","9","4","18","0","2","0","6","0","2","0","2","6","246"),
        ("Management (n=122)",       "45","52","11","10","2","2","0","0","0","0","0","0","0","122"),
        ("Spec. Business (n=158)",   "41","14","53","26","8","0","0","10","0","2","0","2","2","158"),
        ("Marketing (n=116)",        "17","2","9","78","2","0","0","4","2","0","0","0","2","116"),
        ("Writing/Comm (n=48)",      "0","2","4","2","28","5","0","3","0","2","0","0","2","48"),
        ("Humanities (n=70)",        "2","0","0","4","6","52","0","2","0","2","2","0","0","70"),
        ("Social Sciences (n=156)",  "8","6","6","6","4","5","4","93","2","10","7","4","1","156"),
        ("Pre-Prof/Undecl. (n=196)", "29","18","12","19","12","1","2","35","8","14","14","10","22","196"),
        ("Education (n=26)",         "0","0","0","0","0","0","1","5","14","2","2","0","2","26"),
        ("Health Sciences (n=190)",  "0","4","0","12","2","2","0","18","4","130","14","2","2","190"),
        ("Natural Sciences (n=176)", "18","4","2","4","0","0","0","12","2","18","108","6","2","176"),
        ("Engineering (n=214)",      "12","9","0","1","0","2","6","2","0","6","6","164","6","214"),
        ("Computing/Data (n=132)",   "17","4","1","0","2","0","0","8","0","0","0","8","92","132"),
        ("Col. Total",               "386","124","102","180","66","71","13","198","32","188","153","198","139","1,850"),
    ]

    tbl = doc.add_table(rows=0, cols=15)
    tbl.style = "Table Grid"
    # Header row
    hrow = tbl.add_row()
    header_labels = ["Incoming \\ Graduating",
                     "Acct/Fin","Mgmt","Spec\nBus","Mktg",
                     "Writing","Hum.","Social","Pre\nProf",
                     "Educ","Health","Nat\nSci","Eng.","Comp\nData","Total"]
    col_w = [1.35] + [0.38]*13 + [0.38]
    for i, (lbl, w) in enumerate(zip(header_labels, col_w)):
        cell = hrow.cells[i]
        shade_cell(cell, "2F6DB3")
        run = cell.paragraphs[0].add_run(lbl)
        run.bold = True
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        cell.paragraphs[0].paragraph_format.space_after  = Pt(0)

    for ridx, row_data in enumerate(raw_rows):
        row = tbl.add_row()
        is_total = ridx == len(raw_rows) - 1
        fill = "E8EEF7" if is_total else ("F2F4F8" if ridx % 2 == 0 else "FFFFFF")
        for cidx, val in enumerate(row_data):
            cell = row.cells[cidx]
            shade_cell(cell, fill)
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(7.5)
            run.bold = (cidx == 0 or is_total)
            run.font.color.rgb = NAVY if (cidx == 0 or is_total) else GRAY_TEXT
            cell.paragraphs[0].paragraph_format.space_before = Pt(0)
            cell.paragraphs[0].paragraph_format.space_after  = Pt(0)

    doc.add_paragraph()

    add_h2(doc, "Notable Switching Patterns (Major Groups)")
    add_body(doc,
        "Key flows where students left their incoming group at notable rates:"
    )
    add_two_col_table(doc, [
        ("Management → Accounting/Finance/Ops",
         "36.9% of Management students (45 of 122) graduated in Accounting/Finance/Ops. "
         "The largest single cross-group flow in the dataset."),
        ("Specialized Business → Accounting/Finance/Ops",
         "25.9% (41 of 158) — second largest outflow. Specialized Business has the lowest "
         "retention of any declared group at 33.5%."),
        ("Pre-Prof/Undeclared → Social Sciences",
         "17.9% (35 of 196) — the top landing group for undeclared students, followed by "
         "Accounting/Finance/Ops (14.8%) and Computing/Data (11.2%)."),
        ("Natural Sciences → Health Sciences",
         "10.2% (18 of 176) — a meaningful cross-flow reflecting pre-health pathways."),
        ("Health Sciences → Social Sciences",
         "9.5% (18 of 190) — students from health-related majors who re-declared into "
         "Social Sciences (e.g., Psychology)."),
        ("Social Sciences → Health Sciences",
         "6.4% (10 of 156) — reverse flow; Psychology/Social Work students re-declaring "
         "into health-focused majors."),
    ], header=("Flow", "Detail"), col_widths=(2.3, 4.2))

    add_h2(doc, "College-Level Retention Rates")
    add_body(doc,
        "At the college level, retention rates are higher because the definition of "
        "'staying' is broader — a student who switches between two business majors "
        "still counts as retained in OCB."
    )
    add_two_col_table(doc, [
        ("Opus College of Business (OCB)",          "91.3%  —  586 of 642 students  |  incoming n=642"),
        ("School of Engineering (SOE)",              "76.6%  —  164 of 214 students  |  incoming n=214"),
        ("College of Arts and Sciences (CAS)",       "75.3%  —  438 of 582 students  |  incoming n=582"),
        ("Morrison Family College of Health (COH)",  "68.4%  —  130 of 190 students  |  incoming n=190"),
        ("School of Education (EDUC)",               "53.8%  —  14 of 26 students    |  incoming n=26"),
        ("Other / Undeclared",                       "N/A    —  196 students dispersed: 43.9% → CAS, 39.8% → OCB"),
    ], header=("College", "Retention"), col_widths=(3.2, 3.3))

    add_h2(doc, "College-Level Transition Counts (Incoming → Graduating)")
    add_three_col_table(doc, [
        ("OCB → CAS",   "46 students (7.2%)",  "Second most common college switch"),
        ("CAS → OCB",   "90 students (15.5%)", "Largest cross-college flow in the dataset"),
        ("CAS → COH",   "32 students (5.5%)",  "CAS students moving into health sciences"),
        ("COH → CAS",   "38 students (20.0%)", "Health students switching to arts/sciences"),
        ("SOE → OCB",   "22 students (10.3%)", "Engineering students moving to business"),
        ("SOE → CAS",   "22 students (10.3%)", "Engineering students moving to arts/sciences (tie with SOE→OCB)"),
        ("OTHER → CAS", "86 students (43.9%)", "Largest single destination for undeclared students"),
        ("OTHER → OCB", "78 students (39.8%)", "Second destination for undeclared students"),
    ], header=("Flow", "Count", "Note"), col_widths=(1.5, 1.8, 3.2))

    add_h2(doc, "Source Files for This Analysis")
    add_two_col_table(doc, [
        ("analysis_outputs/mobility_tables/mobility_transition_tables.xlsx",
         "Main workbook — 10 sheets: joint counts, raw counts, row %, column %, and "
         "long-format tables for both major-group and college levels."),
        ("analysis_outputs/mobility_tables/mobility_transition_tables_polished.xlsx",
         "Formatted version of the same tables, suitable for direct inclusion in reports."),
        ("analysis_outputs/mobility_tables/*.csv",
         "Individual CSV exports of each sheet for use in Python/R analysis."),
        ("build_transition_tables.py (project root)",
         "Script that generated these tables from Class_of_2026_CPF_vs_SP26.csv + major.csv."),
    ], header=("File", "Contents"), col_widths=(3.0, 3.5))

    # ─────────────────────────────────────────────────────────────────────────
    # 11. NOTES FOR FUTURE STUDENTS
    # ─────────────────────────────────────────────────────────────────────────
    add_h1(doc, "11. Notes for Future Students")
    add_body(doc,
        "If you are picking this project up for the first time, here is the recommended "
        "reading order before touching any code:"
    )
    add_bullet(doc, "1. Open site/index.html in a browser to see what the final product looks like.")
    add_bullet(doc, "2. Read PIPELINE_OVERVIEW.md — understand what each of the 10 steps does.")
    add_bullet(doc, "3. Read DATA_CLEANING_DECISIONS.md — understand why the data looks the way it does.")
    add_bullet(doc, "4. Read VISUALIZATION_RULES.md — understand the color and layout system.")
    add_bullet(doc, "5. Only then: install Circos, update data files, and run python run_pipeline.py.")

    add_h2(doc, "The most common mistakes to avoid")
    add_bullet(doc, "Do not open the HTML files from a restricted network share — "
               "browsers may block local SVG loading. Open from a local drive.")
    add_bullet(doc, "Do not skip the CIRCOS_ROOT environment variable step. Every Circos "
               "render call will fail silently if it cannot find the executable.")
    add_bullet(doc, "When you add new clusters, you must update THREE places: "
               "the remap script (Steps 2 or 3), the HTML file (arc display names array), "
               "and dashboard_docs_data.js (methodology text). Updating only one will "
               "cause mismatches.")
    add_bullet(doc, "The dashboard_docs_data.js file is NOT generated by any script. "
               "It must be updated manually whenever cluster names or graduate counts change.")
    add_bullet(doc, "After any SVG regeneration, bump the cache-buster query strings in "
               "the HTML img src tags (e.g., ?v=20260507 → ?v=20270515). Otherwise "
               "browsers will keep showing old cached SVGs.")

    add_h2(doc, "Where to find archived decisions")
    add_body(doc,
        "The full rationale for every data and visualization decision is documented in the "
        "Markdown files in this package. The original development workspace "
        "(circos_refresh_v2/) contains experiment folders (order_exp_v1 through v7), "
        "comparison HTML files, and early prototype dashboards that show the evolution "
        "of the project. Those are NOT part of this reproducible package but are available "
        "in the original project repository if you need to understand why a particular "
        "approach was tried and discarded."
    )

    # ─────────────────────────────────────────────────────────────────────────
    # Save
    # ─────────────────────────────────────────────────────────────────────────
    doc.save(str(OUTPUT_PATH))
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
